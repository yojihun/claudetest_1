from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Mm


ROOT = Path("/Users/jihoonkim/Desktop/vibe-coding/HIM_exam")
OUTPUT_DIR = ROOT / "artifacts"
OUTPUT_PATH = OUTPUT_DIR / "병원통계_쉽게_공부하는_책.docx"
PAGE_DIR = OUTPUT_DIR / "hospital_statistics_book_pages"
FONT_PATH = Path("/System/Library/Fonts/AppleSDGothicNeo.ttc")

PAGE_WIDTH = 1240
PAGE_HEIGHT = 1754
MARGIN_X = 96
MARGIN_TOP = 110
MARGIN_BOTTOM = 110
BODY_LINE_GAP = 18

NAVY = (18, 36, 58)
TEAL = (0, 132, 116)
BLUE = (40, 90, 210)
GRAY = (100, 112, 132)
LIGHT = (242, 247, 250)
WHITE = (255, 255, 255)


@dataclass
class Block:
    kind: str
    text: str


def font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size)


FONTS = {
    "cover_title": font(54),
    "cover_subtitle": font(28),
    "section_kicker": font(24),
    "h1": font(38),
    "h2": font(28),
    "body": font(31),
    "body_bold": font(31),
    "small": font(20),
}


def content_blocks() -> list[Block]:
    return [
        Block("cover_title", "병원통계, 쉽게 잡는 공부책"),
        Block("cover_subtitle", "보건의료정보관리사 시험 대비\n통계 기초부터 실전 계산까지 완벽하게 정리"),
        Block("body", "이 책은 병원통계가 특히 어렵게 느껴지는 학습자를 위해 만들어진 특화 교재이다. "
             "단순 공식을 기계적으로 외우기보다, 용어의 근본적인 뜻을 파악하고 분모와 분자를 구별하는 눈을 기르도록 돕는다."),
        Block("bullet", "Chapter I. 병원통계 기초와 용어 사전: 기초 지식, 병상, 재원일수, 사망, 외래·응급, 산과, 부검, 역학 및 생정·출산 지표의 쉬운 뜻과 공식을 총망라한다."),
        Block("bullet", "Chapter II. 계산 방법 실전 훈련: 비율형, 평균형, 회전형, 그리고 역학·출산력형으로 공식을 분류하고 실전 문제 해결을 위한 단계별 루틴을 익힌다."),
        Block("bullet", "Chapter III. 빈출 문제 자가 점검: 국가시험에서 자주 출제되는 핵심 질문과 상세 답변 15선을 통해 실전 능력을 최종 확인한다."),
        Block("h2", "효율적인 이 책의 학습 가이드"),
        Block("number", "1단계: 용어 사전을 보며 각 지표의 '분모'와 '분자'에 들어갈 대상이 누군지 한국어로 설명해 본다."),
        Block("number", "2단계: 계산 훈련 챕터에서 비율형, 평균형, 회전형의 특징과 자주 하는 실수(함정)를 숙지한다."),
        Block("number", "3단계: 직접 연습장에 예제 숫자를 넣어서 풀어보고 자가 점검 퀴즈로 약점을 보완한다."),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER I"),
        Block("h1", "병원통계 기초와 용어 사전"),
        Block("body", "병원통계는 숫자를 다루기 전에 공식에 쓰이는 재원(입원 중)과 퇴원(입원 완료)의 기준을 명확히 하는 것에서 시작한다. "
             "용어가 비슷하더라도 그 지표가 측정하려는 실제 운영 목적이나 대상이 무엇인지를 아는 것이 핵심이다."),
        Block("h2", "1. 통계의 기본 개념 (비, 비율, 율)"),
        Block("bullet", "비 (Ratio): 서로 독립적인 두 수량의 비를 뜻한다. (예: 남녀비 = 남성 환자 수 ÷ 여성 환자 수)"),
        Block("bullet", "비율 (Proportion): 전체 중에서 특정 부분이 차지하는 크기다. 분자가 분모에 포함된다. (예: 전체 환자 중 소아 환자 비율)"),
        Block("bullet", "율 (Rate): 시간의 변화에 따라 어떤 사건이 일어난 빈도를 뜻한다. (예: 조사망률, 감염률)"),
        Block("h2", "2. 척도와 대표값, 산포도"),
        Block("term", "척도의 종류: 명명척도(분류), 서열척도(순서), 등간척도(간격 일정, 절대0 없음), 비율척도(절대0 존재, 사칙연산 가능)"),
        Block("term", "대표값: 자료를 대표하는 값으로 산술평균, 중앙값(크기 순서대로 나열했을 때 가운데 값), 최빈값(가장 빈도가 높은 값)이 있다."),
        Block("term", "산포도: 자료가 대표값 주위에 얼마나 흩어져 있는가를 나타내는 척도로, 분산과 표준편차가 대표적이다."),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER I"),
        Block("h1", "병상 및 재원일수 관련 지표"),
        Block("h2", "3. 병상 지표의 쉬운 뜻과 공식"),
        Block("term", "병상이용률 (Bed Occupancy Rate): 준비된 병상이 일정 기간 동안 실제로 얼마나 활용되었는지를 나타내는 지표다. 공식은 (기간 중 총 환자일수 ÷ (가동병상수 × 기간일수)) × 100 (%)"),
        Block("term", "병상회전율 (Bed Turnover Rate): 일정 기간 동안 가동병상 1개당 환자가 몇 번이나 거쳐갔는가(퇴원했는가)를 나타낸다. 공식은 기간 중 총 퇴원환자수 (사망 포함) ÷ 가동병상수 (회)"),
        Block("term", "병상회전간격 (Bed Turnover Interval): 한 환자가 퇴원하고 다음 환자가 들어올 때까지 병상이 비어 있는 평균 날짜다. 공식은 ((가동병상수 × 기간일수) - 총 환자일수) ÷ 퇴원환자수 (일)"),
        Block("h2", "4. 재원일수 지표의 쉬운 뜻과 공식"),
        Block("term", "일일평균재원환자수 (Average Daily Census): 특정 기간 동안 하루 평균 입원해 있었던 환자의 머릿수다. 공식은 기간 중 총 환자일수 ÷ 기간일수 (명)"),
        Block("term", "평균재원일수 (Average Length of Stay): 퇴원한 환자 1명이 입원해서 퇴원할 때까지 평균 며칠이 걸렸는지를 뜻한다. 공식은 퇴원환자의 총 재원일수 ÷ 퇴원환자수 (일)"),
        Block("term", "환자일수 (Inpatient Service Days): 병원에 하루 동안 입원해 있던 환자들의 일수를 모두 누적한 합계다. (예: 3명이 4일씩 입원했다면 환자일수는 12일)"),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER I"),
        Block("h1", "사망 및 부검 관련 지표"),
        Block("h2", "5. 사망 관련 지표 (분모·분자 구분의 핵심)"),
        Block("term", "조사망률 (Gross Mortality Rate): 병원 전체 퇴원환자 중 사망자의 비율이다. 공식은 (기간 중 총 사망수 ÷ 총 퇴원환자수(사망 포함)) × 100 (%)"),
        Block("term", "순사망률 (Net Mortality Rate): 불가항력적인 입원 48시간 미만 사망을 제외한 사망률이다. 공식은 (48시간 이상 사망수 ÷ (총 퇴원환자수 - 48시간 미만 사망수)) × 100 (%)"),
        Block("term", "신생아사망률 (Newborn Mortality Rate): 신생아실 퇴원아 중 사망아의 비율이다. 공식은 (신생아실 총 사망수 ÷ 신생아실 총 퇴원수) × 100 (%)"),
        Block("term", "수술사망률 (Postoperative Mortality Rate): 수술 후 10일 이내에 발생한 사망률이다. 공식은 (수술 후 10일 이내 사망수 ÷ 총 수술환자수) × 100 (%)"),
        Block("term", "모성사망률 (Maternal Mortality Rate): 임신·분만·산욕기 합병증으로 사망한 비율이다. 공식은 (모성 사망수 ÷ 모성 퇴원환자수) × 100 (%)"),
        Block("h2", "6. 부검 관련 지표"),
        Block("term", "조부검률 (Gross Autopsy Rate): 병원 내 총 사망자 중 병원에서 부검이 집도된 비율이다. 공식은 (총 부검수 ÷ 총 사망자수) × 100 (%)"),
        Block("term", "순부검률 (Net Autopsy Rate): 법의학 사건 등으로 부검이 불가능했던 사망자를 제외한 실제 부검률이다. 공식은 (총 부검수 ÷ (총 사망자수 - 부검제외시신수)) × 100 (%)"),
        Block("term", "병원부검률 (Hospital Autopsy Rate): 진료받던 원외 사망자가 병원으로 이송되어 부검한 케이스까지 합산하여 구한다. 공식은 ((원내부검 + 원외이송부검) ÷ (원내사망 + 원외이송사망)) × 100 (%)"),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER I"),
        Block("h1", "외래·응급·역학 및 출산력 지표"),
        Block("h2", "7. 외래 및 응급환자 통계"),
        Block("term", "외래 신환자 비율: 전체 외래 방문 횟수 중 최초 방문 신규 환자의 비율이다. 공식은 (신환자수 ÷ 외래환자 연인원) × 100 (%)"),
        Block("term", "응급실 사망률: 응급실 총 내원 환자 중 사망하여 퇴원한 비율이다. 공식은 (응급실 사망자수 ÷ 응급실 총 내원환자수) × 100 (%)"),
        Block("term", "응급실 입원율: 응급실 내원객 중 즉각 일반 병상으로 입원 완료된 비율이다. 공식은 (응급실 경유 입원환자수 ÷ 응급실 총 내원환자수) × 100 (%)"),
        Block("h2", "8. 역학 통계 및 생정·출산력 지표"),
        Block("term", "발생률 (Incidence Rate): 특정 기간 건강한 인구 중 질병이 새로 발생한 속도이다. 공식은 (신규 환자수 ÷ 위험노출 연앙인구) × 1,000 (‰)"),
        Block("term", "유병률 (Prevalence Rate): 특정 시점 인구 중 질병을 앓고 있는 모든 환자 비율이다. 공식은 (전체 환자수 ÷ 총 인구수) × 1,000 (‰)"),
        Block("term", "치명률 (Case Fatality Rate): 해당 질환자 중에서 그 병으로 사망한 비율이다. 공식은 (질병 사망자수 ÷ 질병 총 환자수) × 100 (%)"),
        Block("term", "상대위험도 (Relative Risk): 위험 요인 노출군의 발생률과 비노출군의 발생률의 비를 나타낸다. 공식은 노출군 발생률 ÷ 비노출군 발생률 (배)"),
        Block("term", "조출생률 (Crude Birth Rate): 전체 연앙인구 1,000명당 태어난 활생아수 비율이다. 공식은 (연간 총 출생아수 ÷ 연앙인구) × 1,000 (‰)"),
        Block("term", "일반출산율 (General Fertility Rate): 15-49세 가임여성 인구 1,000명당 태어난 활생아수 비율이다. 공식은 (연간 총 출생아수 ÷ 15~49세 여성 인구) × 1,000 (‰)"),
        Block("term", "합계출산율 (Total Fertility Rate): 한 여성이 평생(15~49세) 낳을 것으로 예상되는 평균 자녀수다. 단위는 (명)이다."),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER II"),
        Block("h1", "계산 방법 실전 훈련"),
        Block("body", "병원통계 계산 문제를 풀 때는 공식에 바로 대입하려 하지 말고 다음 3단계 루틴을 적용해야 실수를 방지할 수 있다. "
             "첫째, 구하라는 값의 '단위'를 확인한다. 둘째, '분모'를 구성할 대상을 찾는다. 셋째, 분자에 올릴 타깃 사건 수를 파악한다."),
        Block("h2", "1. 비율/백분율형 문제 적용 (곱하기 100 필수)"),
        Block("bullet", "병상이용률 분모 함정 피하기: 분모는 '가동병상수 × 기간일수'다. 30일인지, 31일인지 달력 일수를 반드시 확인한다."),
        Block("bullet", "사망률 분모 함정 피하기: 분모인 '총 퇴원환자수'에는 사망퇴원자가 이미 포함되어 있음을 인지한다."),
        Block("formula", "순사망률 = (48시간 이상 사망수 ÷ (총 퇴원환자수 - 48시간 미만 사망수)) × 100"),
        Block("example", "실전 문제: 총 퇴원환자 400명, 총 사망 12명(48시간 미만 4명, 48시간 이상 8명)일 때 순사망률은?\n계산: (8 ÷ (400 - 4)) × 100 = (8 ÷ 396) × 100 ≒ 2.02%"),
        Block("h2", "2. 평균형 문제 적용 (나누기 기준 명확화)"),
        Block("bullet", "평균재원일수 분모는 항상 '퇴원환자수'다. 입원환자나 재원환자수로 나누지 않는다."),
        Block("formula", "평균재원일수 = 퇴원환자의 총 재원일수 ÷ 퇴원환자수"),
        Block("example", "실전 문제: 10월 한 달 퇴원환자 200명, 이들의 총 재원일수 합이 1,600일일 때 평균재원일수는?\n계산: 1,600일 ÷ 200명 = 8일"),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER II"),
        Block("h1", "회전형 및 역학·출산력형 계산"),
        Block("h2", "3. 회전형 문제 적용 (병상의 회전 주기 평가)"),
        Block("bullet", "병상회전율은 가동병상 1개당 평균 몇 명의 환자가 거쳐갔는가를 측정하며 단위는 '회'이다."),
        Block("bullet", "병상회전간격은 병상이 비어 있던 평균 날짜를 계산하며 단위는 '일'이다. 분자는 '총 가능 병상일수 - 실제 환자일수'다."),
        Block("formula", "병상회전간격 = ((가동병상수 × 기간일수) - 환자일수) ÷ 퇴원환자수"),
        Block("example", "실전 문제: 가동병상 100개, 30일간 총 환자일수 2,700일, 퇴원환자 150명일 때 회전간격은?\n계산: ((100 × 30) - 2,700) ÷ 150 = (3,000 - 2,700) ÷ 150 = 300 ÷ 150 = 2일"),
        Block("h2", "4. 역학 및 출산력 지표 계산 (천분율 단위 유의)"),
        Block("bullet", "일반출산율과 발생률, 유병률, 조출생률 등은 백분율(%)이 아니라 천분율(‰)을 쓰므로 반드시 마지막에 1,000을 곱한다."),
        Block("formula", "일반출산율 = (연간 총 활생아수 ÷ 15~49세 가임 여성 연앙인구) × 1,000"),
        Block("example", "실전 문제: 가임기 여성 10,000명 중 연간 150명의 활생아가 출생했다. 일반출산율은?\n계산: (150 ÷ 10,000) × 1,000 = 15‰"),
        Block("pagebreak", ""),

        Block("kicker", "CHAPTER III"),
        Block("h1", "자가 점검 (Q&A 15선)"),
        Block("question", "1. 환자일수(Inpatient Service Days)와 퇴원환자 총 재원일수(Discharge Days)의 결정적인 차이는 무엇인가?"),
        Block("answer", "환자일수는 기간 중 매일 입원해 있던 환자들의 일수 합으로 '현재 재원' 관점이고, 퇴원환자 총 재원일수는 해당 기간에 퇴원 완료된 환자들의 입원 기간만을 합산한 것으로 '퇴원' 관점이다."),
        Block("question", "2. 가동병상수가 200개인 병원의 6월(30일) 중 일일평균재원환자수가 160명일 때, 이 달의 병상이용률은 몇 %인가?"),
        Block("answer", "일일평균재원환자수가 160명이므로 30일간 총 환자일수는 160 × 30 = 4,800일이다. 병상이용률 = (4,800 ÷ (200 × 30)) × 100 = 80% 이다."),
        Block("question", "3. 병상회전율과 평균재원일수, 병상이용률의 상관관계는 어떻게 되는가?"),
        Block("answer", "평균재원일수가 짧을수록, 그리고 병상이용률이 높을수록 병상회전율은 높아진다."),
        Block("question", "4. 순사망률을 계산할 때 분모와 분자에서 공통적으로 빼주어야 하는 대상은 누구인가?"),
        Block("answer", "입원 후 48시간 이내에 사망한 환자 수이다."),
        Block("question", "5. 임상 모성사망률을 구할 때 분모에 들어가야 하는 집단은 인구학적 모성사망비와 어떻게 다른가?"),
        Block("answer", "임상 모성사망률의 분모는 '모성 퇴원환자수(사망 포함)'이며, 인구학적 모성사망비의 분모는 '해당 연도의 총 출생아 수(보통 10만 명당)'로 계산한다."),
        Block("question", "6. 태아사망률(사산율)을 구할 때 분모에 대입해야 하는 산식은 무엇인가?"),
        Block("answer", "살아서 태어난 신생아수 (총 분만수) + 태아사망수(사산아 수)이다."),
        Block("question", "7. 수술사망률을 구할 때 수술 후 며칠 이내의 사망을 분자에 합산하는가?"),
        Block("answer", "수술 후 10일 이내에 발생한 사망을 분자에 포함한다."),
        Block("question", "8. 병상회전간격이 음수(-)로 계산되는 특수한 상황은 어떤 경우인가?"),
        Block("answer", "실제 환자일수가 이론상 가동병상일수보다 많을 때 발생하며, 이는 임시 병상을 추가 가동하여 초과 입원을 받았음을 시사한다."),
        Block("question", "9. 분산과 표준편차가 0이라는 것은 자료들이 어떻게 분포해 있다는 뜻인가?"),
        Block("answer", "모든 자료의 수치가 평균값과 완전히 동일하여 흩어짐이 전혀 없다는 뜻이다."),
        Block("question", "10. 병원 통계 문제를 풀기 직전 검증해야 하는 3단계 최종 확인 요소는 무엇인가?"),
        Block("answer", "첫째, 소수점 반올림 기준. 둘째, 계산하라는 지표의 단위(%, 일, 명, 회, ‰). 셋째, 분모에서 48시간 미만 등 제외 대상이 있는지 여부이다."),
        Block("question", "11. 조사부검률과 순부검률의 계산 공식에서 분모의 차이는 무엇인가?"),
        Block("answer", "조부검률의 분모는 '총 사망자수'이지만, 순부검률의 분모는 '총 사망자수 - 부검 불가능 사망수(법의학 사건 등)'이다."),
        Block("question", "12. 발생률(Incidence)과 유병률(Prevalence)의 개념적 차이는 무엇인가?"),
        Block("answer", "발생률은 '신규 발생 환자'의 속도를 보며 분모에 기존 유병자는 제외하지만, 유병률은 특정 시점에 '기존 및 신규 질환 전체 환자'의 규모를 측정한다."),
        Block("question", "13. 조출생률(CBR)과 일반출산율(GFR)의 분모 차이는 무엇인가?"),
        Block("answer", "조출생률은 전체 '연앙인구'를 분모로 쓰지만, 일반출산율은 실제 출산이 가능한 '15-49세 가임여성 연앙인구'만을 분모로 쓴다."),
        Block("question", "14. 상대위험도(RR)가 1이라는 것의 의미는 무엇인가?"),
        Block("answer", "위험 요인에 노출된 군과 노출되지 않은 군의 질병 발생률이 동일하여, 해당 요인과 질병 발생 간에 연관성이 전혀 없음을 뜻한다."),
        Block("question", "15. 치명률(Case Fatality Rate)의 분모는 다른 사망률들과 어떻게 다른가?"),
        Block("answer", "다른 사망률은 전체 퇴원환자나 인구를 분모로 삼지만, 치명률은 오직 '해당 질병에 걸린 전체 환자수(이환자수)'만을 분모로 삼는다."),
        Block("body", "마무리 조언: 병원통계는 단순 연산 공식의 조합이 아니라 환자와 병상 자원의 유기적 흐름을 숫자로 표현한 과목이다. 각 공식의 정의를 우리말로 설명할 수 있다면 국가시험 계산은 결코 어렵지 않다."),
    ]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, text_font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    pieces = text.split("\n")
    lines: list[str] = []
    for piece in pieces:
        if not piece:
            lines.append("")
            continue
        current = ""
        for token in piece.split(" "):
            candidate = token if not current else f"{current} {token}"
            if draw.textlength(candidate, font=text_font) <= width:
                current = candidate
                continue
            if current:
                lines.append(current)
                current = token
                continue
            broken = wrap(piece, width=max(1, width // max(1, text_font.size // 2)))
            lines.extend(broken)
            current = ""
            break
        if current:
            lines.append(current)
    return lines or [""]


def block_layout(draw: ImageDraw.ImageDraw, block: Block, usable_width: int) -> tuple[list[str], ImageFont.FreeTypeFont, tuple[int, int, int], int, int]:
    if block.kind == "cover_title":
        return wrap_text(draw, block.text, FONTS["cover_title"], usable_width), FONTS["cover_title"], NAVY, 64, 18
    if block.kind == "cover_subtitle":
        return wrap_text(draw, block.text, FONTS["cover_subtitle"], usable_width), FONTS["cover_subtitle"], TEAL, 46, 12
    if block.kind == "kicker":
        return wrap_text(draw, block.text, FONTS["section_kicker"], usable_width), FONTS["section_kicker"], TEAL, 34, 8
    if block.kind == "h1":
        return wrap_text(draw, block.text, FONTS["h1"], usable_width), FONTS["h1"], NAVY, 36, 8
    if block.kind == "h2":
        return wrap_text(draw, block.text, FONTS["h2"], usable_width), FONTS["h2"], BLUE, 28, 8
    if block.kind == "formula":
        return wrap_text(draw, block.text, FONTS["body_bold"], usable_width - 28), FONTS["body_bold"], NAVY, 18, 10
    prefix_map = {
        "bullet": "• ",
        "number": "",
        "term": "",
        "example": "예시 ",
        "question": "",
        "answer": "정리 ",
    }
    if block.kind in prefix_map:
        text = block.text
        if block.kind == "number":
            text = block.text
        elif block.kind == "term":
            head, tail = block.text.split(": ", 1)
            text = f"{head}\n{tail}"
        elif block.kind == "question":
            text = f"Q. {block.text.split('. ', 1)[1] if '. ' in block.text else block.text}"
        elif block.kind == "answer":
            text = f"A. {block.text}"
        else:
            text = f"{prefix_map[block.kind]}{block.text}"
        return wrap_text(draw, text, FONTS["body"], usable_width - 18), FONTS["body"], NAVY, 16, 8
    return wrap_text(draw, block.text, FONTS["body"], usable_width), FONTS["body"], NAVY, 18, 8


def measure_block(draw: ImageDraw.ImageDraw, block: Block, usable_width: int) -> int:
    lines, text_font, _, before, after = block_layout(draw, block, usable_width)
    line_height = text_font.size + BODY_LINE_GAP
    return before + len(lines) * line_height + after


def draw_block(draw: ImageDraw.ImageDraw, block: Block, x: int, y: int, usable_width: int) -> int:
    lines, text_font, color, before, after = block_layout(draw, block, usable_width)
    y += before
    line_height = text_font.size + BODY_LINE_GAP

    if block.kind == "formula":
        box_top = y - 8
        box_height = len(lines) * line_height + 24
        draw.rounded_rectangle((x, box_top, x + usable_width, box_top + box_height), radius=24, fill=LIGHT, outline=(210, 222, 232), width=2)
        text_x = x + 16
        for idx, line in enumerate(lines):
            draw.text((text_x, y + idx * line_height), line, font=text_font, fill=color)
        return box_top + box_height + after

    if block.kind in {"question", "answer"}:
        box_color = (246, 250, 255) if block.kind == "question" else (240, 248, 243)
        border = (208, 220, 235) if block.kind == "question" else (198, 224, 208)
        box_top = y - 8
        box_height = len(lines) * line_height + 24
        draw.rounded_rectangle((x, box_top, x + usable_width, box_top + box_height), radius=20, fill=box_color, outline=border, width=2)
        for idx, line in enumerate(lines):
            draw.text((x + 16, y + idx * line_height), line, font=text_font, fill=color)
        return box_top + box_height + after

    for idx, line in enumerate(lines):
        draw.text((x, y + idx * line_height), line, font=text_font, fill=color)
    return y + len(lines) * line_height + after


def render_pages(blocks: list[Block]) -> list[Path]:
    PAGE_DIR.mkdir(parents=True, exist_ok=True)
    for old in PAGE_DIR.glob("page-*.png"):
        old.unlink()

    probe = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), WHITE)
    probe_draw = ImageDraw.Draw(probe)
    usable_width = PAGE_WIDTH - MARGIN_X * 2
    usable_height = PAGE_HEIGHT - MARGIN_TOP - MARGIN_BOTTOM

    pages: list[list[Block]] = []
    current: list[Block] = []
    current_height = 0

    for block in blocks:
        if block.kind == "pagebreak":
            if current:
                pages.append(current)
                current = []
                current_height = 0
            continue

        needed = measure_block(probe_draw, block, usable_width)
        if current and current_height + needed > usable_height:
            pages.append(current)
            current = []
            current_height = 0
        current.append(block)
        current_height += needed

    if current:
        pages.append(current)

    page_paths: list[Path] = []
    for index, page_blocks in enumerate(pages, start=1):
        img = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), WHITE)
        draw = ImageDraw.Draw(img)
        y = MARGIN_TOP

        if index > 1:
            draw.line((MARGIN_X, 70, PAGE_WIDTH - MARGIN_X, 70), fill=(228, 236, 243), width=2)
            draw.text((MARGIN_X, 34), "병원통계, 쉽게 잡는 공부책", font=FONTS["small"], fill=GRAY)

        for block in page_blocks:
            y = draw_block(draw, block, MARGIN_X, y, usable_width)

        page_label = f"{index}"
        label_width = draw.textlength(page_label, font=FONTS["small"])
        draw.text((PAGE_WIDTH - MARGIN_X - label_width, PAGE_HEIGHT - 70), page_label, font=FONTS["small"], fill=GRAY)

        page_path = PAGE_DIR / f"page-{index}.png"
        img.save(page_path, quality=95)
        page_paths.append(page_path)

    return page_paths


def build_docx(page_paths: list[Path]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)

    image_width = Mm(210)
    for page_path in page_paths:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Mm(0)
        paragraph.paragraph_format.space_after = Mm(0)
        run = paragraph.add_run()
        run.add_picture(str(page_path), width=image_width)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


def main() -> None:
    blocks = content_blocks()
    page_paths = render_pages(blocks)
    build_docx(page_paths)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
